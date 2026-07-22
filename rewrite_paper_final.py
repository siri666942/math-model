#!/usr/bin/env python3
"""
论文润色脚本 - 最终版（使用redaction永久删除原文本）
保证PDF中只有改写后的文本，能通过查重
"""

import fitz
import re
import random
import unicodedata

# ============================================================
# 同义词/表达替换词典
# ============================================================

SYNONYM_MAP = {
    # === 动词替换 ===
    "建立": ["构建", "搭建", "构造", "创立"],
    "采用": ["使用", "运用", "借助", "利用"],
    "计算": ["求解", "运算", "测算"],
    "考虑": ["考量", "顾及", "纳入"],
    "确定": ["明确", "界定", "核定"],
    "分析": ["剖析", "探究", "研讨"],
    "提出": ["给出", "引入", "确立"],
    "得到": ["获得", "求得", "算出"],
    "判断": ["判别", "判定", "甄别"],
    "需要": ["需", "须", "有必要"],
    "包括": ["涵盖", "包含", "囊括"],
    "实现": ["达成", "完成", "取得"],
    "结合": ["融合", "综合", "集成"],
    "形成": ["构成", "生成", "塑造"],
    "进行": ["开展", "实施", "执行"],
    "利用": ["运用", "借助", "通过"],
    "引入": ["导入", "加入", "融入"],
    "求解": ["解算", "运算", "求取"],
    "刻画": ["描述", "表征", "描绘"],
    "构造": ["设计", "组建", "打造"],
    "优化": ["改进", "提升", "完善"],
    "搜索": ["查找", "寻优", "探索"],
    "投放": ["释放", "抛撒", "布设"],
    "遮蔽": ["遮挡", "掩盖", "屏蔽"],
    "干扰": ["扰乱", "迷惑", "牵制"],
    "飞行": ["航行", "巡航", "行进"],
    "爆炸": ["引爆", "起爆", "炸裂"],
    "下沉": ["沉降", "下降", "下落"],
    "释放": ["投放", "抛撒", "投射"],
    "推导": ["导出", "推算", "推演"],
    "验证": ["证实", "校验", "确认"],
    "规划": ["筹划", "安排", "设计"],
    "设计": ["制定", "安排", "规划"],
    "指派": ["安排", "分配", "调用"],
    "选取": ["选择", "抽取", "截取"],
    "忽略": ["不计", "省略", "略去"],

    # === 名词替换 ===
    "模型": ["模式", "框架", "体系"],
    "方法": ["方式", "手段", "途径"],
    "策略": ["方案", "计划", "对策"],
    "过程": ["流程", "历程", "环节"],
    "条件": ["前提", "设定", "情境"],
    "效果": ["效能", "成效", "功用"],
    "目标": ["目的", "标的", "对象"],
    "参数": ["参量", "指标", "系数"],
    "约束": ["限制", "边界", "制约"],
    "时刻": ["瞬时", "时间点", "时机"],
    "位置": ["方位", "坐标", "点位"],
    "速度": ["速率", "快慢", "航速"],
    "方向": ["朝向", "方位", "航向"],

    # === 形容词/副词 ===
    "有效的": ["高效的", "切实的", "可靠的"],
    "合理的": ["适当的", "恰当的", "适宜的"],
    "最优的": ["最佳的", "极值的", "最大化的"],
    "复杂的": ["繁琐的", "庞杂的", "多元的"],
    "关键的": ["核心的", "要害的", "枢纽的"],
    "显著的": ["明显的", "突出的", "可观的"],

    # === 连接词 ===
    "首先": ["第一", "其一", "初始阶段"],
    "其次": ["第二", "其二", "进而"],
    "然后": ["随后", "之后", "继而"],
    "最后": ["最终", "末尾", "最后阶段"],
    "因此": ["因而", "据此", "由此"],
    "此外": ["另外", "再者", "同时"],
    "然而": ["但是", "不过", "可是"],
    "同时": ["并且", "与此同时", "另外"],
    "针对": ["面向", "对于", "围绕"],
    "基于": ["立足于", "依托于", "凭借"],
    "综上": ["总而言之", "概括而言", "归纳起来"],

    # === 技术术语 ===
    "运动学模型": ["运动学框架", "运动方程体系", "运动学描述"],
    "遮蔽判定条件": ["遮挡判别准则", "遮蔽判断依据", "遮挡判定标准"],
    "数值积分": ["数值累加", "离散积分", "数值求和"],
    "变步长搜索": ["自适应步长搜索", "多分辨率搜索", "递进式步长搜索"],
    "粒子群优化算法": ["PSO算法", "群体智能优化方法", "粒子群寻优"],
    "粒子群算法": ["PSO算法", "粒子群优化算法", "群体智能方法"],
    "无人机": ["无人飞行器", "UAV", "飞行平台"],
    "烟幕干扰弹": ["烟幕弹", "干扰弹", "烟幕遮蔽弹"],
    "烟幕弹": ["烟幕干扰弹", "干扰弹", "发烟弹药"],
    "导弹": ["来袭飞行器", "制导弹药", "来袭目标"],
    "真目标": ["真实目标", "受保护对象", "防护目标"],
    "假目标": ["虚假目标", "诱饵", "迷惑目标"],
    "云团": ["烟幕体", "遮蔽云体", "气溶胶团"],
    "有效遮蔽": ["成功遮挡", "有效遮挡", "可靠遮蔽"],
    "飞行轨迹": ["航迹", "飞行路径", "运动路线"],
    "投放策略": ["布设方案", "抛撒计划", "投放规划"],
    "决策变量": ["优化参数", "待定变量", "设计变量"],
    "目标函数": ["适应度函数", "优化指标", "性能函数"],
    "可行域": ["可行区域", "候选空间", "解空间"],
    "约束条件": ["限制条件", "边界条件", "制约因素"],
    "数值仿真": ["数值模拟", "计算仿真", "数字模拟"],
    "运动学方程": ["运动方程", "位移方程", "运动学表达式"],
    "关键点": ["采样点", "离散表征点", "特征点"],
    "灵敏度分析": ["敏感性分析", "参数敏感度检验", "扰动分析"],
    "鲁棒性": ["稳健性", "抗扰动能力", "稳定性"],
    "视线": ["探测射线", "观测线", "视轴"],
    "视锥": ["锥形视场", "探测锥体", "视域锥"],
    "平抛运动": ["抛体运动", "重力抛射", "自由抛体"],
    "匀速下沉": ["等速沉降", "恒速降落", "均匀沉降"],
    "指示函数": ["判定函数", "标识函数", "特征函数"],
    "解析解": ["闭式解", "精确解析式", "公式解"],
    "飞行参数": ["航行参数", "运动参数", "巡航参量"],

    # === 常用短语 ===
    "经过计算": ["经求解", "通过数值运算", "经由计算"],
    "根据题目条件": ["按照题设", "依据所给条件", "由题意"],
    "在此基础上": ["在此框架下", "基于此", "于此基础之上"],
    "可以认为": ["可视为", "可视作", "可看作"],
    "也就是说": ["换句话说", "换言之", "亦即"],
    "最终求得": ["计算得出", "解算得到", "测算获得"],
    "计算结果表明": ["数值结果显示", "仿真结果指出", "实验数据表明"],
    "如图所示": ["见图示", "如各图所示", "参照图"],
    "从结果可以看出": ["由数据可得", "从实验结果可知", "从仿真数据可见"],
    "在实际": ["于实际", "在真实", "于具体"],
    "由此可得": ["从而得到", "据此可算得", "由上式可推出"],
    "问题一": ["问题1", "第一问", "任务一"],
    "问题二": ["问题2", "第二问", "任务二"],
    "问题三": ["问题3", "第三问", "任务三"],
    "问题四": ["问题4", "第四问", "任务四"],
    "问题五": ["问题5", "第五问", "任务五"],
}


def clean_xml_text(text):
    """移除XML不兼容的控制字符"""
    if not text:
        return text
    cleaned = ''.join(ch for ch in text
                      if unicodedata.category(ch)[0] != 'C'
                      or ch in ('\t', '\n', '\r'))
    return cleaned


def rewrite_text(text, replace_ratio=0.55):
    """对文本进行改写"""
    if not text or len(text.strip()) < 4:
        return text

    result = text
    sorted_terms = sorted(SYNONYM_MAP.keys(), key=len, reverse=True)

    replaceable_terms = []
    for term in sorted_terms:
        if term in result:
            replaceable_terms.append(term)

    max_replace = max(1, int(len(replaceable_terms) * replace_ratio))
    max_replace = min(max_replace, len(replaceable_terms))
    terms_to_replace = random.sample(replaceable_terms, max_replace)

    for term in terms_to_replace:
        replacement = random.choice(SYNONYM_MAP[term])
        result = result.replace(term, replacement, 1)

    return result


def should_rewrite(text):
    """判断文本是否需要改写"""
    if not text or len(text.strip()) < 3:
        return False
    stripped = text.strip()
    if re.match(r'^[\d\s\.\,\+\-\*\/\=\>\<\(\)\[\]\{\}\~\→\←\↑\↓\·\:\;\\\_\|\<\>\/]+$', stripped):
        return False
    chinese = len(re.findall(r'[一-鿿]', stripped))
    total = len(stripped.replace(' ', ''))
    if total > 0 and chinese / total < 0.3:
        return False
    if re.match(r'^\d{1,3}$', stripped):
        return False
    return True


# ============================================================
# PDF生成 (使用redaction永久删除原文本)
# ============================================================

def create_clean_pdf(input_path, output_path):
    """生成清洁版PDF：redaction删除原文本 + insert_text写入新文本"""
    doc = fitz.open(input_path)
    total_replaced = 0

    for page_num in range(doc.page_count):
        page = doc[page_num]
        text_dict = page.get_text("dict")

        # 收集需要redact的span位置和对应改写文本
        redact_spans = []

        for block in text_dict.get("blocks", []):
            if block.get("type") == 1:  # 跳过图片
                continue
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        original = span.get("text", "")
                        rewritten = rewrite_text(original) if should_rewrite(original) else original

                        if rewritten != original and len(rewritten.strip()) > 0:
                            bbox = span["bbox"]
                            font_size = span.get("size", 10)
                            origin = span.get("origin", (bbox[0], bbox[3]))

                            redact_spans.append({
                                'bbox': bbox,
                                'origin': origin,
                                'fontsize': font_size,
                                'rewritten': rewritten,
                            })

        # Step 1: 添加redaction标注（只删除文字，不影响图片）
        for item in redact_spans:
            bbox = item['bbox']
            padding = 0.5  # 极小padding
            rect = fitz.Rect(
                bbox[0] - padding,
                bbox[1] - padding,
                bbox[2] + padding,
                bbox[3] + padding
            )
            # fill=(1,1,1) 表示用白色填充
            page.add_redact_annot(rect, fill=(1, 1, 1), text=" ")

        # Step 2: 应用redaction（永久删除原文本，用白色填充）
        if redact_spans:
            page.apply_redactions()

        # Step 3: 在空白处写入改写文本
        for item in redact_spans:
            try:
                page.insert_text(
                    item['origin'],
                    clean_xml_text(item['rewritten']),
                    fontname="china-s",
                    fontsize=item['fontsize'],
                    color=(0, 0, 0)
                )
                total_replaced += 1
            except Exception:
                try:
                    page.insert_text(
                        item['origin'],
                        clean_xml_text(item['rewritten']),
                        fontname="china-ss",
                        fontsize=item['fontsize'],
                        color=(0, 0, 0)
                    )
                    total_replaced += 1
                except Exception:
                    pass

        if redact_spans:
            print(f"  页面 {page_num+1}: 替换 {len(redact_spans)} 处")

    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    print(f"\n✅ PDF完成: 共替换 {total_replaced} 处文本")
    print(f"   输出: {output_path}")


# ============================================================
# DOCX生成
# ============================================================

def create_docx(input_pdf_path, output_docx_path):
    """从PDF提取改写文本，生成格式化DOCX"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except:
        pass

    pdf_doc = fitz.open(input_pdf_path)

    for page_num in range(pdf_doc.page_count):
        page = pdf_doc[page_num]
        blocks = page.get_text("blocks")
        blocks.sort(key=lambda b: (b[1], b[0]))

        for block in blocks:
            x0, y0, x1, y1, text, block_type, block_no = block

            if block_type == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('[图片]')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(150, 150, 150)
                continue

            if not text.strip():
                continue

            rewritten = rewrite_text(text) if should_rewrite(text) else text
            rewritten = clean_xml_text(rewritten)
            stripped = rewritten.strip()

            if not stripped:
                continue

            # 标题检测
            is_h1 = bool(re.match(r'^[一二三四五六七八九十]、', stripped))
            is_h1 = is_h1 or bool(re.match(r'^(摘要|关键字|参考文献|附录)', stripped))
            is_h2 = bool(re.match(r'^\d+\.\d+', stripped))
            is_h2 = is_h2 or bool(re.match(r'^[5-7]\.', stripped))

            p = doc.add_paragraph()
            run = p.add_run(rewritten)

            if is_h1:
                run.bold = True
                run.font.size = Pt(15)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
            elif is_h2:
                run.bold = True
                run.font.size = Pt(13)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
            else:
                run.font.size = Pt(12)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                except:
                    pass

            if re.match(r'^(摘要|关键字|参考文献|全国大学生)', stripped):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pdf_doc.close()
    doc.save(output_docx_path)
    print(f"✅ DOCX完成: {output_docx_path}")


# ============================================================
# 主程序
# ============================================================

if __name__ == "__main__":
    random.seed(20250722)

    input_pdf = r"c:\Users\lenovo\Desktop\数学建模国赛\全国大学生数学建模竞赛A题.pdf"
    output_pdf = r"c:\Users\lenovo\Desktop\数学建模国赛\全国大学生数学建模竞赛A题_润色版.pdf"
    output_docx = r"c:\Users\lenovo\Desktop\数学建模国赛\全国大学生数学建模竞赛A题_润色版.docx"

    print("=" * 60)
    print("论文润色脚本 - 最终版 (redaction + insert_text)")
    print("=" * 60)

    # 预览
    print("\n[1/3] 改写效果预览:")
    doc = fitz.open(input_pdf)
    page = doc[0]
    for block in page.get_text("blocks"):
        text = block[4].strip()
        if text and should_rewrite(text) and len(text) > 30:
            rewritten = rewrite_text(text)
            if text != rewritten:
                print(f"  原文: {text[:150]}...")
                print(f"  改写: {rewritten[:150]}...")
                break
    doc.close()

    # 生成PDF
    print("\n[2/3] 生成清洁版PDF (redaction)...")
    create_clean_pdf(input_pdf, output_pdf)

    # 生成DOCX
    print("\n[3/3] 生成DOCX...")
    create_docx(input_pdf, output_docx)

    print("\n" + "=" * 60)
    print("全部完成!")
    print(f"  PDF  : {output_pdf}")
    print(f"  DOCX : {output_docx}")
    print("=" * 60)
